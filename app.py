from flask import Flask, render_template, request, flash, redirect, url_for, session, send_from_directory,jsonify, send_file, Response, make_response
from io import BytesIO
# from xhtml2pdf import pisa
from bs4 import BeautifulSoup
import pdfkit
import psycopg2
import os
from flask_sqlalchemy import SQLAlchemy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import linear_kernel
import glob
from docx import Document
from PyPDF2 import PdfReader
from PyPDF2 import PdfFileReader
import fitz
import re
from flask_wtf import FlaskForm
from wtforms import RadioField, StringField, PasswordField, SubmitField, validators
from werkzeug.utils import secure_filename
import base64
from wtforms import BooleanField, MultipleFileField, StringField, TextAreaField, FileField, SubmitField
from wtforms.validators import DataRequired, InputRequired
from flask_wtf.file import FileField, FileAllowed

from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
from pdfme import PDF




# Flask app configuration
app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Change this to a secure secret key
app.config['SQLALCHEMY_DATABASE_URI'] = 'postgresql://postgres:root@localhost:8080/postgres'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'

db = SQLAlchemy(app)





class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(255), nullable=False)
    type = db.Column(db.String(255), nullable=False)
    description = db.Column(db.String(255), nullable=True)
    pdf_path = db.Column(db.String(255), nullable=False)
    approval_status = db.Column(db.String(255), nullable=False)

class User(db.Model):
    __tablename__ = 'users1'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(20), unique=True, nullable=False)
    password = db.Column(db.String(60), nullable=False)
    role=db.Column(db.String(20),nullable=False)
    email=db.Column(db.String(30),unique=True)
    manager_id =db.Column(db.Integer)
    manager_name=db.Column(db.String(20))

class LoginForm(FlaskForm):
    username = StringField('Username', [validators.InputRequired()])
    password = PasswordField('Password', [validators.InputRequired()])
    #role = RadioField('Role', choices=[('manager', 'Manager'), ('employee', 'Employee')], validators=[validators.InputRequired()])
    submit = SubmitField('Login')

admin_credentials = {'username': 'admin', 'password': 'admin123'}

class Answer(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    content = db.Column(db.Text, nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('users1.id'), nullable=False)
    user = db.relationship('User', backref=db.backref('answers', lazy=True))
    question_id = db.Column(db.Integer, db.ForeignKey('question.id'), nullable=False)
    edited = db.Column(db.Boolean, default=False)
    edited_at = db.Column(db.DateTime, nullable=True)

class AnswerForm(FlaskForm):
    content = TextAreaField('Your Answer', validators=[InputRequired()])
    submit = SubmitField('Post Answer')
    edited = BooleanField('edited', default=False)

# Route for the login page
@app.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        username = form.username.data
        password = form.password.data

        if username == admin_credentials['username'] and password == admin_credentials['password']:
            return redirect(url_for('admin_dashboard'))
        #role=form.role.data
        user = User.query.filter_by(username=username).first()

        if user and user.password == password:
            #flash('Login successful', 'success')
            
                session['username'] = username
               # current_user = session.get('username')
                return redirect(url_for('dashboard',username=username))
            
        else:
            #flash('Login unsuccessful. Please check your credentials.', 'danger')
            return redirect(url_for('login'))

    return render_template('login.html', form=form)


@app.route('/changepassword<username>', methods=['GET', 'POST'])
def changepassword(username):
    if request.method == 'POST':
        old_password = request.form.get('old_password')
        new_password = request.form.get('new_password')
        confirm_password = request.form.get('confirm_password')

        #current_user = session.get('username')
        
        # Validate current password
        user = User.query.filter_by(username=username).first()

        if user and user.password == old_password:
            if new_password == confirm_password:
                # Update password
                user.password = new_password
                db.session.commit()
                flash('Password successfully changed', 'success')
                return redirect(url_for('dashboard', username=username))
                
            else:
                flash('New password and confirm password do not match', 'error')
        else:
            flash('Incorrect current password', 'error')

    return render_template('changepassword.html',username=username)

# Route for the dashboard
@app.route('/dashboard')
def dashboard():
    if 'username' in session:
        username = session['username']
        return render_template('dashboard.html', username=username)
    else:
        return redirect(url_for('login'))

# Route for logout
@app.route('/logout')
def logout():
    session.pop('username', None)
    session.pop('role', None)
    return redirect(url_for('login'))


@app.route('/add_post_doc', methods=['POST', 'GET'])
def add_post_doc():
    if request.method == 'POST':
        option = request.form.get('option')

        doc_title = request.form['doc_title']
        brief_description = request.form['brief_description']
        
        doc_type = request.form['type']

        if option == 'write':
            content = request.form['content']
            # Create a directory to store written documents if not exists
            if doc_type == 'Learning':
                write_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'Learning')

            elif doc_type == 'Training':
                write_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'Training')

            elif doc_type == 'KT':
                write_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'KT')

            elif doc_type == 'ProjectSpecific':
                write_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'ProjectSpecific')

            elif doc_type == 'UserSpecific':
                write_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'UserSpecific')

            else:
                write_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'Others')
            
            os.makedirs(write_dir, exist_ok=True)
            # Save the document as a text file
            file_path = os.path.join(write_dir, f"{doc_title}.txt")
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(f"{content}")

            # Save document information in the database
            new_document = Document(title=doc_title, type=doc_type, description=brief_description, pdf_path=file_path, approval_status='Pending')
            db.session.add(new_document)
            db.session.commit()

            return redirect(url_for('add_post_doc'))

    return render_template('add_post_doc.html')

def get_folder_content(path):
    content = []
    with os.scandir(path) as entries:
        for entry in entries:
            info = {
                'name': entry.name,
                'is_dir': entry.is_dir(),
                'path': entry.path.replace(os.sep, '/'),
            }
            if entry.is_dir():
                info['content'] = get_folder_content(entry.path)
            content.append(info)
    return content

# @app.route('/display_files', methods=['POST','GET'])
# def display_files():
#     directory = 'C:/Users/7000035069/Desktop/PF_SOCIALS/uploads'
#     folder_content = get_folder_content(directory)
#     if request.method == 'POST':
#         query = request.form.get('search_query')
#         results = search_documents(query)
#         return render_template('display_files.html', folder_content=folder_content, get_folder_content=get_folder_content, query=query, results=results)
#     return render_template('display_files.html', folder_content=folder_content, get_folder_content=get_folder_content, query='', results=[])

def split_string(value, separator):
    return value.split(separator)

app.jinja_env.filters['split'] = split_string

@app.route('/view/<path:filename>')
def view_file(filename):
    file_path = os.path.join(".", filename).replace(os.sep, '/') 
    return render_template('view_file.html', file_path=file_path)

@app.route('/filesold/<path:filename>')
def get_file(filename):
    parts = re.split(r'[\\\/]', filename)
    undesired_parts = [0,1,2,3,4]
    remaining_parts = [part for index, part in enumerate(parts) if index not in undesired_parts]
    result = '/'.join(remaining_parts)
    return send_from_directory('uploads',result)



# Function to read text from different document types
def read_text(file_path):
    _, file_extension = os.path.splitext(file_path)
    if file_extension == '.pdf':
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            return ' '.join(page.extract_text() for page in pdf_reader.pages)
    elif file_extension == '.docx':
        doc = Document(file_path)
        return ' '.join(paragraph.text for paragraph in doc.paragraphs)
    elif file_extension == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            return file.read()
    else:
        return ''  # Unsupported file type

# Function to search documents by name and content
def search_documents(query):
    results = []
    allowed_extensions = ['.pdf', '.txt']

    for root, dirs, files in os.walk('C:/Users/7000035069/Desktop/PF_SOCIALS/uploads'):
        for file in files:
            if file.lower().endswith(tuple(allowed_extensions)):
                file_path = os.path.join(root, file)

                if file.lower().endswith('.pdf'):
                    pdf_document = fitz.open(file_path)
                    for page_num in range(pdf_document.page_count):
                        page = pdf_document[page_num]
                        page_text = page.get_text()
                        if query and txt_content is not None and (query.lower() in page_text.lower() or query.lower() in file.lower()):
                            parts = re.split(r'[\\\/]', file_path)
                            undesired_parts = []
                            remaining_parts = [part for index, part in enumerate(parts) if index not in undesired_parts]
                            result = '/'.join(remaining_parts)
                            results.append(result)
                            break
                    pdf_document.close()

                elif file.lower().endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8') as txt_file:
                        txt_content = txt_file.read()
                        if query and txt_content is not None and (query.lower() in txt_content.lower() or query.lower() in file.lower()):
                            parts = re.split(r'[\\\/]', file_path)
                            undesired_parts = []
                            remaining_parts = [part for index, part in enumerate(parts) if index not in undesired_parts]
                            result = '/'.join(remaining_parts)
                            results.append(result)


    return results

@app.route('/search', methods=['POST'])
def search():
    data = request.get_json()
    query = data.get('query', '')
    results = search_documents(query)
    return jsonify({'results': results})

@app.route('/admin_dashboard', methods=['POST','GET'])
def admin_approval():
    directory = 'C:/Users/7000035069/Desktop/PF_SOCIALS/uploads'
    folder_content = get_folder_content(directory)
    if request.method == 'POST':
        query = request.form.get('search_query')
        results = search_documents(query)
        return render_template('admin_dashboard1.html', folder_content=folder_content, get_folder_content=get_folder_content, query=query, results=results)
    return render_template('admin_dashboard1.html', folder_content=folder_content, get_folder_content=get_folder_content, query='', results=[])

@app.route('/summary', methods=['POST','GET'])
def create_summary():
    return render_template('create_summary.html')

@app.route('/query', methods=['POST','GET'])
def query_page():
    return render_template('query_page.html')

@app.route('/publish', methods=['POST','GET'])
def publish():
    return render_template('publish.html')

ALLOWED_EXTENSIONS = {'txt', 'pdf'}
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/add_post_file', methods=['POST', 'GET'])
def add_post_file():
    if request.method == 'POST':
        option = request.form.get('option')
        doc_title = request.form['doc_title']
        brief_description = request.form['brief_description']
        doc_type = request.form['type']

        

        if option == 'upload':
                uploaded_file = request.files['file']
                # Create a directory to store uploaded files if not exists
                if doc_type == 'Learning':
                    upload_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'Learning')

                elif doc_type == 'Training':
                    upload_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'Training')

                elif doc_type == 'KT':
                    upload_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'KT')

                elif doc_type == 'ProjectSpecific':
                    upload_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'ProjectSpecific')

                elif doc_type == 'UserSpecific':
                    upload_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'UserSpecific')

                else:
                    upload_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'Others')
                
                os.makedirs(upload_dir, exist_ok=True)
                # Save the uploaded file
                if uploaded_file:
                    file_path = os.path.join(upload_dir, uploaded_file.filename)
                    uploaded_file.save(file_path)
                    # Save document information in the database
                    new_document = Document(title=doc_title, description=brief_description, pdf_path=file_path, type=doc_type, approval_status='Pending')
                    db.session.add(new_document)
                    db.session.commit()
                    return redirect(url_for('add_post_file'))
                
                if 'file' not in request.files:
                    return jsonify({'error': 'No file part'})

                if uploaded_file.filename == '':
                    return jsonify({'error': 'No selected file'})

    return render_template('add_post_file.html')

class Question(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(100), nullable=False)
    content = db.Column(db.Text, nullable=False)
    image_data = db.Column(db.ARRAY(db.String), nullable=True)  # Array to store multiple images
    user_id = db.Column(db.Integer, db.ForeignKey('users1.id'), nullable=False)
    user = db.relationship('User', backref=db.backref('questions', lazy=True))
    answers = db.relationship('Answer', backref='question', lazy=True)
 
class QuestionForm(FlaskForm):
    title = StringField('Title', validators=[InputRequired()])
    content = TextAreaField('Content', validators=[InputRequired()])
    images = FileField('Upload Images', validators=[FileAllowed(['jpg', 'jpeg', 'png', 'gif'])])
    submit = SubmitField('Ask Question')
 
@app.route('/add_post_ask', methods=['GET', 'POST'])
def ask_question():
    form = QuestionForm()
    user = User.query.filter_by(username=session['username']).first()
    if form.validate_on_submit():
        title = form.title.data
        content = form.content.data
        images = request.files.getlist('images')
 
        # Process and store the images
        image_data = []
 
        for image_file in images:
            if image_file:
                image_data.append(base64.b64encode(image_file.read()).decode('utf-8'))
            else:
                image_data.append(None)
        question = Question(title=title, content=content, user=user, image_data=image_data)
        db.session.add(question)
        db.session.commit()
       
 
        #flash('Question posted successfully!', 'success')
        return redirect(url_for('ask_question', username=session['username']))
 
    return render_template('add_post_ask.html', form=form,user=user)

@app.route("/<username>")
def index(username):
    user = User.query.filter_by(username=username).first()
    if user:
        questions = Question.query.all()
        return render_template("index.html", questions=questions, user=user)
    else:
        flash('User not found.', 'danger')
        return redirect(url_for('login'))

@app.route('/add_post_post', methods=['POST','GET'])
def add_post_post():
    return render_template('add_post_post.html')

@app.route('/display_files', methods=['POST','GET'])
def display_files():
    return render_template("display_files.html")

base_folder_path = 'C:/Users/7000035069/Desktop/PF_SOCIALS/uploads'

def get_file_info(file_path):
    return {
        'name': os.path.basename(file_path),
        'path': file_path,
        'isFolder': os.path.isdir(file_path)
    }

def get_files_in_folder(folder_path):
    files = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        files.append(get_file_info(file_path))
    return files

@app.route('/files', methods=['GET'])
def get_root_files():
    try:
        files = get_files_in_folder(base_folder_path)
        return jsonify(files)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/files/<path:subpath>', methods=['GET'])
def get_subfolder_content(subpath):
    try:
        subfolder_path = os.path.join(base_folder_path, subpath)
        files = get_files_in_folder(subfolder_path)
        return jsonify(files)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/file/<path:file_path>', methods=['GET'])
def get_file_content(file_path):
    try:
        if os.path.isfile(file_path):
            
            parts = file_path.split('/')
            undesired_parts = [0,1,2,3,4]
            remaining_parts = [part for index, part in enumerate(parts) if index not in undesired_parts]
            result = '/'.join(remaining_parts)
            print(result)
            with open(result, 'rb') as file:
                content = file.read()
            return Response(content, content_type='application/pdf')
        else:
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
# @app.route('/file-link/<path:file_path>', methods=['GET'])
# def get_file_link(file_path):
#     try:
#         if file_path.endswith('.pdf'):
#             link = f'http://127.0.0.1:5000//{file_path}'
#             return jsonify({'link': link})
#         else:
#             return jsonify({'error': 'Invalid file type'}), 400
#     except Exception as e:
#         return jsonify({'error': str(e)}), 500

# @app.route('/generate_pdf', methods=['POST'])
# def generate_pdf():
#     content = request.form['content']

#     # Create a PDF document
#     pdf_filename = 'output.pdf'
#     with open(pdf_filename, 'wb') as f:
#         c = canvas.Canvas(f)
#         c.drawString(100, 750, content)  # Adjust the position as needed
#         c.save()

#     return send_file(pdf_filename, as_attachment=True)
    


from fpdf import FPDF
# from html2pdf import HTML2PDF

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'PDF Generator', 0, 1, 'C')

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, 'Page %s' % self.page_no(), 0, 0, 'C')


@app.route('/generate_pdf', methods=['POST'])
def generate_pdf():
    content = request.form['content']

    # Create a PDF document
    pdf = PDF()
    pdf.add_page()

    # Set font and size
    pdf.set_font("Arial", size=12)

    # Parse HTML content
    soup = BeautifulSoup(content, 'html.parser')
    formatted_text = soup.get_text("\n", strip=True)

    # Add content to the PDF
    pdf.multi_cell(0, 10, formatted_text)

    # Save the PDF to a file
    pdf_filename = 'output.pdf'
    pdf.output(pdf_filename)

    return send_file(pdf_filename, as_attachment=True)

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        app.run(debug=True)
